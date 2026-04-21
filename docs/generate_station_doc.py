from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT_PATH = Path(__file__).resolve().parent / "Documentacion_Estacion_Terrena_y_Terrena_Web.docx"


def add_title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "Documentacion de la Estacion Terrena y la Terrena Web", "KA'AN ASTRA")

    doc.add_paragraph(
        "Este documento describe el funcionamiento del software de la estacion terrena local, "
        "la arquitectura de la terrena web, el flujo de datos entre ambos sistemas y el mecanismo "
        "de respaldo mediante base de datos. La redaccion se centra en el comportamiento real del "
        "proyecto actual y en la forma en que sus modulos trabajan de manera integrada."
    )

    doc.add_heading("1. Vision general del sistema", level=1)
    doc.add_paragraph(
        "El sistema completo se compone de tres bloques principales: el sistema de recepcion fisica, "
        "la estacion terrena de escritorio y la terrena web. Los modulos de radio LoRa y XBee entregan "
        "sus datos a un microcontrolador ESP receptor. Ese ESP concentra la informacion y la envia por "
        "un unico puerto serial hacia la computadora donde corre la estacion terrena. La estacion terrena "
        "procesa los datos, actualiza sus vistas locales, genera reportes y retransmite la telemetria a "
        "un backend en la nube. Finalmente, la terrena web consulta ese backend y presenta la informacion "
        "en tiempo real a usuarios remotos."
    )
    add_bullet(doc, "Transmisor -> modulos de radio -> ESP receptor -> puerto serial -> estacion terrena local")
    add_bullet(doc, "Estacion terrena local -> Worker API -> base de datos D1 -> terrena web")

    doc.add_heading("2. Como funciona la estacion terrena local", level=1)
    doc.add_paragraph(
        "La estacion terrena local esta construida con Electron. Esto permite combinar una interfaz "
        "de escritorio con acceso al puerto serial, al sistema de archivos y a procesos locales. Su "
        "funcion principal es recibir telemetria, interpretarla, mostrarla y mantener una salida "
        "sincronizada hacia la plataforma web."
    )

    doc.add_heading("2.1 Recepcion de datos", level=2)
    doc.add_paragraph(
        "La recepcion se realiza desde un solo puerto serial, ya que el ESP receptor actua como "
        "concentrador de LoRa y XBee. El software lista los puertos disponibles, permite seleccionar "
        "uno desde la interfaz y abre la comunicacion serial con la configuracion definida. A partir "
        "de ese momento, cada linea que llega al parser es evaluada para determinar si corresponde "
        "a telemetria util."
    )
    add_bullet(doc, "Soporta datos CSV del formato tradicional de la estacion terrena")
    add_bullet(doc, "Soporta mensajes etiquetados como [LORA] y [XBEE]")
    add_bullet(doc, "Soporta mensajes parciales con etiquetas como LAT, LON, TEMP, HUM, ALT, etc.")
    add_bullet(doc, "Ignora lineas que son solo logs o mensajes sin telemetria relevante")

    doc.add_heading("2.2 Parser y normalizacion", level=2)
    doc.add_paragraph(
        "El software utiliza una capa de parsing que transforma distintos formatos de entrada en un "
        "objeto interno comun. Esto permite que el sistema sea compatible tanto con telemetria completa "
        "como con telemetria parcial. Si el receptor solo envia latitud y longitud, la estacion puede "
        "seguir trabajando con esos datos. Si posteriormente se agregan nuevas magnitudes etiquetadas, "
        "el mismo parser las reconoce sin necesidad de rediseñar toda la aplicacion."
    )
    doc.add_paragraph(
        "Los datos normalizados se almacenan en una estructura interna de estado. De esta manera, si llega "
        "telemetria parcial en una muestra, el sistema puede conservar datos previamente validos y actualizar "
        "solo los campos nuevos. Esto es especialmente util durante la fase actual, en la que el receptor puede "
        "transmitir solo posicion, pero en una etapa futura podra transmitir la telemetria completa del CanSat."
    )

    doc.add_heading("2.3 Calculos derivados", level=2)
    doc.add_paragraph(
        "La estacion terrena no solo muestra valores directamente recibidos. Tambien calcula variables derivadas "
        "a partir de la informacion disponible. Entre las mas importantes estan la velocidad horizontal estimada, "
        "la altitud relativa, la velocidad vertical, la aceleracion total y la distancia entre receptor y transmisor."
    )
    add_bullet(doc, "Velocidad horizontal: se estima con el cambio de posicion GPS entre muestras")
    add_bullet(doc, "Altitud relativa y velocidad vertical: se obtienen a partir del filtrado y fusion de sensores")
    add_bullet(doc, "Aceleracion total: se calcula desde los ejes del acelerometro")
    add_bullet(doc, "Distancia receptor-transmisor: se calcula con la formula de Haversine si hay coordenadas TX y RX")

    doc.add_heading("2.4 Visualizacion local", level=2)
    doc.add_paragraph(
        "La interfaz local contiene varias vistas: un dashboard principal, un mapa y un modelo 3D. El dashboard usa "
        "graficas temporales para representar magnitudes como temperatura, humedad, presion, aceleracion, altitud, "
        "velocidad, viento y distancia entre receptor y transmisor. El mapa permite seguir la posicion de la carga y "
        "su trayectoria. El modelo 3D visualiza la orientacion estimada con base en datos giroscopicos."
    )
    doc.add_paragraph(
        "La estacion terrena fue ampliada para incluir una grafica especifica de distancia. Esto permite ver la "
        "evolucion temporal de la separacion entre el transmisor y la terrena receptora, lo cual resulta util para "
        "pruebas de alcance, seguimiento de posicion y verificacion operacional en campo."
    )

    doc.add_heading("2.5 Generacion de reportes", level=2)
    doc.add_paragraph(
        "La aplicacion local genera reportes en Excel y en texto. El Excel conserva la estructura tabular de la "
        "telemetria, mientras que el archivo de texto resume estadisticas, estado de la mision y resultados de interes. "
        "Estos reportes se producen a partir del historial acumulado durante la ejecucion de la sesion."
    )

    doc.add_heading("2.6 Sincronizacion con la nube", level=2)
    doc.add_paragraph(
        "Una vez que la estacion terrena construye el objeto de telemetria final, lo envia por HTTP a un backend "
        "desplegado en Cloudflare Workers. Este envio se realiza de forma asincrona para no bloquear la interfaz local. "
        "Asi, la estacion terrena local actua como puente entre la recepcion fisica y la plataforma remota."
    )

    doc.add_heading("3. Como funciona la terrena web", level=1)
    doc.add_paragraph(
        "La terrena web es una interfaz externa publicada en Cloudflare Pages. Su objetivo es reflejar, en tiempo real, "
        "la informacion que recibe la estacion terrena local, para que pueda ser consultada por mas personas desde un "
        "navegador. No se conecta directamente a los radios ni al ESP. Toda su informacion proviene del backend remoto."
    )

    doc.add_heading("3.1 Interfaz y estructura", level=2)
    doc.add_paragraph(
        "La terrena web reproduce la logica visual de la estacion terrena de escritorio. Cuenta con un dashboard principal "
        "de graficas, una subpestaña de mapa y una subpestaña de modelo 3D. Esto permite mantener una experiencia de uso "
        "similar entre la interfaz local y la remota."
    )
    add_bullet(doc, "Dashboard: muestra graficas de temperatura, humedad, presion, aceleracion, altitud, viento, velocidad y distancia")
    add_bullet(doc, "Mapa: muestra trayectoria, coordenadas, distancia y controles de centrado")
    add_bullet(doc, "Modelo 3D: muestra la orientacion del objeto usando datos giroscopicos")

    doc.add_heading("3.2 Flujo de consulta", level=2)
    doc.add_paragraph(
        "La terrena web consulta periodicamente el backend usando endpoints como /api/latest, /api/recent y /api/report. "
        "Con esa informacion actualiza las graficas, el mapa, la posicion actual y el modelo 3D. Tambien dispone de un "
        "boton para descargar reportes con nombre basado en fecha y hora local del sistema del usuario."
    )

    doc.add_heading("3.3 Mapa y modelo 3D", level=2)
    doc.add_paragraph(
        "El mapa en la web usa Leaflet y representa la trayectoria del transmisor, la posicion actual y la informacion "
        "asociada de coordenadas y distancia. El modelo 3D utiliza Three.js. En la implementacion actual puede usar un "
        "modelo real en formato GLB, lo que permite reemplazar facilmente la representacion generada por geometria simple "
        "por un archivo 3D del objeto real."
    )

    doc.add_heading("3.4 Rol de la web dentro del sistema", level=2)
    doc.add_paragraph(
        "La terrena web no sustituye a la estacion local. Su funcion es ser una ventana remota de observacion y consulta. "
        "El punto de adquisicion real sigue siendo la estacion terrena de escritorio, mientras que la web actua como "
        "plataforma de visualizacion compartida para jueces, operadores remotos o publico autorizado."
    )

    doc.add_heading("4. Sistema de respaldo y almacenamiento en base de datos", level=1)
    doc.add_paragraph(
        "Para evitar depender solo de la memoria local de la estacion terrena, el sistema incorpora un mecanismo de respaldo "
        "remoto mediante Cloudflare D1. Cada muestra de telemetria enviada por la estacion de escritorio es almacenada en la "
        "base de datos a traves del Worker API. Esto permite conservar historico, asegurar consistencia entre sesiones y "
        "ofrecer datos confiables a la terrena web."
    )

    doc.add_heading("4.1 Objetivo del respaldo", level=2)
    add_bullet(doc, "Evitar perdida de informacion en caso de cierre de la app local")
    add_bullet(doc, "Dar consistencia a la terrena web entre multiples consultas")
    add_bullet(doc, "Conservar historico de la mision")
    add_bullet(doc, "Permitir generar reportes descargables desde la nube")
    add_bullet(doc, "Separar la recepcion fisica del almacenamiento persistente")

    doc.add_heading("4.2 Flujo del respaldo", level=2)
    doc.add_paragraph(
        "El flujo del respaldo es el siguiente: la estacion local recibe una muestra, la procesa, genera el objeto de "
        "telemetria y lo envia al Worker mediante POST /api/telemetry. El Worker valida la estructura, normaliza los campos "
        "y los inserta en la base D1. Posteriormente, la terrena web consulta latest, recent o report para obtener la "
        "informacion desde la base de datos."
    )

    doc.add_heading("4.3 Que datos se almacenan", level=2)
    doc.add_paragraph(
        "La base de datos almacena tanto telemetria principal como metadatos utiles para la operacion. Entre los campos se "
        "encuentran hora, velocidad, temperatura, humedad, presion, aceleraciones, giroscopio, magnetometro, altitud, "
        "latitud, longitud, velocidad horizontal, velocidad vertical, altitud relativa, desacople, canal de origen y "
        "distancia al receptor. Tambien se almacena un timestamp UTC de recepcion para mantener trazabilidad temporal."
    )

    doc.add_heading("4.4 Ventajas operativas", level=2)
    doc.add_paragraph(
        "El respaldo en D1 mejora la robustez del sistema porque ya no depende de memoria efimera del backend. De esta forma, "
        "la web siempre consulta una fuente persistente y consistente. Ademas, el historico queda disponible para analisis, "
        "descarga y auditoria posterior."
    )

    doc.add_heading("5. Integracion completa del sistema", level=1)
    doc.add_paragraph(
        "El comportamiento global del sistema puede resumirse en la siguiente secuencia: el transmisor envia datos, los radios "
        "LoRa y/o XBee los entregan al ESP receptor, el ESP reenvia esa informacion por un solo serial a la estacion terrena, "
        "la estacion la interpreta y la visualiza localmente, despues la retransmite a la nube, el Worker la almacena en D1 y "
        "la terrena web la consulta para representarla de forma remota. Este esquema permite una operacion local robusta y al "
        "mismo tiempo una observacion remota en tiempo real."
    )

    doc.add_heading("6. Estado actual y crecimiento futuro", level=1)
    doc.add_paragraph(
        "En la etapa actual, el sistema ya puede operar con telemetria parcial, especialmente con coordenadas GPS de transmisor "
        "y receptor. Esto hace posible seguir posicion, trayectoria, velocidad estimada y distancia entre nodos. A futuro, el "
        "mismo esquema permitira incorporar la telemetria completa contemplada por la estacion terrena, sin necesidad de "
        "reformular la arquitectura, ya que el parser y la base de datos fueron preparados para crecer de forma incremental."
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
